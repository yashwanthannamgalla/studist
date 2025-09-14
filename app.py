import os
import json
import re
import random
from datetime import datetime
from flask import (
    Flask, render_template, request, redirect, send_file,
    send_from_directory, session, jsonify, url_for
)
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
from io import BytesIO
from docx import Document
import spacy

import openai

# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'your_secret_key')  # Change for production!

# Load spaCy model
nlp = spacy.load('en_core_web_sm', disable=['parser', 'ner'])

# Allowed file extensions
ALLOWED_DOC_EXTS = {'pdf', 'png', 'jpg', 'jpeg', 'docx'}
ALLOWED_HANDWRITING_EXTS = {'pdf', 'png', 'jpg', 'jpeg'}

# Upload folders setup
UPLOAD_FOLDER = 'uploads'
NOTES_FOLDER = 'user_notes'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(NOTES_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32 MB

# --- Helper Functions ---
def ensure_user_folder(username: str) -> str:
    user_folder = os.path.join(app.config['UPLOAD_FOLDER'], username)
    os.makedirs(user_folder, exist_ok=True)
    return user_folder

def allowed(filename: str, exts: set) -> bool:
    return ('.' in filename) and (filename.rsplit('.', 1)[1].lower() in exts)

def load_json(path: str, default):
    if os.path.exists(path):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return default
    return default

def save_json(path: str, data) -> None:
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4, ensure_ascii=False)

def to_spotify_embed(url: str) -> str:
    m = re.match(r'https?://open\.spotify\.com/(playlist|track|album|artist|show|episode)/([A-Za-z0-9]+)', url)
    if not m:
        return "https://open.spotify.com/embed/playlist/37i9dQZF1DXcBWIGoYBM5M"
    kind, sid = m.group(1), m.group(2)
    return f"https://open.spotify.com/embed/{kind}/{sid}"

def user_notes_path(username: str) -> str:
    return os.path.join(NOTES_FOLDER, f"{username}_notes.json")

def load_notes(username: str):
    return load_json(user_notes_path(username), [])

def save_notes(username: str, notes):
    save_json(user_notes_path(username), notes)

# ---------- Root/Login/Signup/Logout ----------
@app.route('/', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        username = request.form['username'].strip()
        password = request.form['password']
        users = load_json('users.json', [])
        for user in users:
            if user.get('username') == username and user.get('password') == password:
                session['username'] = username
                return redirect(f"/dashboard?user={username}")
        error = "❌ Invalid username or password. Please try again."
    return render_template('login.html', error=error)

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    error = None
    if request.method == 'POST':
        username = request.form['username'].strip()
        password = request.form['password']
        users = load_json('users.json', [])
        if any(u.get('username') == username for u in users):
            error = "⚠️ Username already exists. Try another one."
            return render_template('signup.html', error=error)
        users.append({'username': username, 'password': password})
        save_json('users.json', users)
        return redirect('/')
    return render_template('signup.html', error=error)

@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect('/')

@app.route('/favicon.ico')
def favicon():
    return ('', 204)

# ---------- Dashboard ----------
@app.route('/dashboard')
def dashboard():
    username = request.args.get('user') or session.get('username')
    if not username:
        return redirect('/login')
    reminders_db = load_json('reminders.json', {})
    reminders_raw = reminders_db.get(username, [])
    reminders = [r['title'] if isinstance(r, dict) else str(r) for r in reminders_raw]
    user_folder = ensure_user_folder(username)
    uploaded_files = [f for f in sorted(os.listdir(user_folder)) if os.path.isfile(os.path.join(user_folder, f))]
    sp_conf = load_json('spotify.json', {})
    spotify_url = sp_conf.get(username, "https://open.spotify.com/embed/playlist/37i9dQZF1DXcBWIGoYBM5M")
    daily_quotes = [
        "The secret of getting ahead is getting started. — Mark Twain",
        "Don’t watch the clock; do what it does. Keep going. — Sam Levenson",
        # ... more quotes ...
        "Work hard in silence, let your success be the noise."
    ]
    daily_quote = random.choice(daily_quotes)
    return render_template(
        'dashboard.html',
        username=username,
        reminders=reminders,
        files=uploaded_files,
        spotify_url=spotify_url,
        daily_quote=daily_quote
    )

# ---------- Notifications (Assignments/Reminders) ----------
def get_user_assignments(username: str):
    db = load_json('assignments.json', {})
    data = db.get(username, [])
    return [a for a in data if isinstance(a, dict)]

def get_user_reminders(username: str):
    db = load_json('reminders.json', {})
    data = db.get(username, [])
    out = []
    for r in data:
        if isinstance(r, dict):
            out.append(r)
        else:
            out.append({'title': str(r), 'date': '', 'time': ''})
    return out

@app.route("/notifications_data")
def notifications_data():
    username = request.args.get('user') or session.get('username')
    if not username:
        return jsonify([])
    assignments = get_user_assignments(username)
    reminders = get_user_reminders(username)
    notifications = []
    for a in assignments:
        title = a.get('subject') or a.get('title') or 'Assignment'
        due = a.get('due_date', '')
        text = f"Assignment: {title}" + (f" (Due {due})" if due else "")
        notifications.append({
            "type": "assignment",
            "text": text,
            "link": f"/assignments?user={username}"
        })
    now = datetime.now()
    for r in reminders:
        title = r.get('title') or 'Reminder'
        date = r.get('date') or ''
        time_ = r.get('time') or ''
        show = True
        if date:
            try:
                dt = datetime.strptime(f"{date} {time_ or '00:00'}", "%Y-%m-%d %H:%M")
                show = now >= dt
            except Exception:
                show = True
        if show:
            notifications.append({
                "type": "reminder",
                "text": f"Reminder: {title}",
                "link": f"/reminders?user={username}"
            })
    return jsonify(notifications)

# ---------- Uploads ----------
@app.route('/upload', methods=['GET', 'POST'])
def upload():
    username = request.args.get('user') or session.get('username')
    if not username:
        return redirect('/')
    user_folder = ensure_user_folder(username)
    if request.method == 'POST':
        if 'file' not in request.files:
            return "No file part"
        file = request.files['file']
        if file.filename == '':
            return "No selected file"
        if allowed(file.filename, ALLOWED_DOC_EXTS):
            file.save(os.path.join(user_folder, secure_filename(file.filename)))
            return redirect(f"/upload?user={username}")
        return "File type not allowed"
    files = [f for f in os.listdir(user_folder) if os.path.isfile(os.path.join(user_folder, f))]
    return render_template('upload.html', username=username, files=files)

@app.route('/uploads/<username>/<filename>')
def uploaded_file(username, filename):
    user_folder = ensure_user_folder(username)
    return send_from_directory(user_folder, filename)

@app.route('/delete-file', methods=['POST'])
def delete_file():
    username = request.form['username']
    filename = request.form['filename']
    path = os.path.join(ensure_user_folder(username), filename)
    if os.path.exists(path):
        os.remove(path)
    return redirect(f"/upload?user={username}")

# ---------- Bookmarks ----------
@app.route('/save-bookmark', methods=['POST'])
def save_bookmark():
    username = request.json['username']
    filename = request.json['filename']
    position = request.json['position']
    data = load_json('bookmarks.json', {})
    data.setdefault(username, {})[filename] = position
    save_json('bookmarks.json', data)
    return jsonify({'message': 'Bookmark saved'})

@app.route('/load-bookmark/<username>/<filename>')
def load_bookmark(username, filename):
    data = load_json('bookmarks.json', {})
    return jsonify({'position': data.get(username, {}).get(filename, 0)})

# ---------- Reminders ----------
@app.route('/reminders')
def view_reminders():
    username = request.args.get('user') or session.get('username')
    if not username:
        return redirect('/')
    data = load_json('reminders.json', {})
    if request.args.get('json'):
        return jsonify(data.get(username, []))
    return render_template('reminder.html', username=username, reminders=data.get(username, []))

@app.route('/add-reminder', methods=['GET', 'POST'])
def add_reminder():
    username = request.args.get('user') or session.get('username')
    if not username:
        return redirect('/')
    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        date = request.form.get('date', '')
        time_ = request.form.get('time', '')
        data = load_json('reminders.json', {})
        data.setdefault(username, []).append({'title': title, 'date': date, 'time': time_})
        save_json('reminders.json', data)
        return redirect(f"/reminders?user={username}")
    return render_template('add_reminder.html', username=username)

@app.route('/delete-reminder', methods=['POST'])
def delete_reminder():
    username = request.form.get('username')
    title = request.form.get('title')
    data = load_json('reminders.json', {})
    data[username] = [r for r in data.get(username, []) if r.get('title') != title]
    save_json('reminders.json', data)
    return redirect(f"/reminders?user={username}")

# ---------- Assignments ----------
@app.route('/assignments')
def assignments():
    username = request.args.get('user') or session.get('username')
    if not username:
        return redirect('/')
    data = load_json('assignments.json', {})
    return render_template('assignments.html', username=username, assignments=data.get(username, []))

@app.route('/add-assignment', methods=['GET', 'POST'])
def add_assignment():
    username = request.args.get('user') or session.get('username')
    if not username:
        return redirect('/')
    if request.method == 'POST':
        subject = request.form['subject'].strip()
        due_date = request.form['due_date']
        description = request.form['description'].strip()
        data = load_json('assignments.json', {})
        data.setdefault(username, []).append({
            'subject': subject,
            'due_date': due_date,
            'description': description,
            'completed': False
        })
        save_json('assignments.json', data)
        return redirect(f"/assignments?user={username}")
    return render_template('add_assignment.html', username=username)

@app.route('/update-assignment', methods=['POST'])
def update_assignment():
    username = request.json['username']
    subject = request.json['subject']
    status = bool(request.json['completed'])
    data = load_json('assignments.json', {})
    for a in data.get(username, []):
        if a.get('subject') == subject:
            a['completed'] = status
            break
    save_json('assignments.json', data)
    return jsonify({'message': 'Assignment status updated'})

@app.route('/delete-assignment', methods=['POST'])
def delete_assignment():
    username = request.json['username']
    subject = request.json['subject']
    data = load_json('assignments.json', {})
    data[username] = [a for a in data.get(username, []) if a.get('subject') != subject]
    save_json('assignments.json', data)
    return jsonify({'message': 'Assignment deleted successfully'})

# ---------- Handwriting Assignment Generator ----------
openai.api_key = os.environ.get('OPENAI_API_KEY')
def allowed_handwriting_file(filename):
    return allowed(filename, ALLOWED_HANDWRITING_EXTS)

def generate_assignment_text(topic):
    prompt = f"Write a detailed assignment on the topic: {topic}."
    try:
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=prompt,
            max_tokens=500,
            temperature=0.7,
        )
        return response.choices[0].text.strip()
    except Exception as e:
        return f"Error generating assignment: {e}"

@app.route('/upload-handwriting', methods=['POST'])
def upload_handwriting():
    topic = request.form.get('topic')
    file = request.files.get('file')
    if not topic or not file:
        return jsonify({'error': 'Missing topic or file'}), 400
    if not allowed(file.filename, ALLOWED_DOC_EXTS):
        return jsonify({'error': 'File type not allowed'}), 400
    filename = secure_filename(file.filename)
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(path)
    # Optional: process file for handwriting/extract first page
    assignment_text = generate_assignment_text(topic)
    doc = Document()
    doc.add_heading(topic, level=0)
    doc.add_paragraph(assignment_text)
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return send_file(
        byte_io,
        as_attachment=True,
        download_name=f"{topic}_assignment.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# ---------- Generate Assignment Form ----------
def generate_assignment_text(topic):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant who writes detailed assignments."},
                {"role": "user", "content": f"Write a detailed assignment on the topic: {topic}."}
            ],
            max_tokens=500,
            temperature=0.7,
        )
        return response['choices'][0]['message']['content'].strip()
    except Exception as e:
        return f"Error generating assignment: {e}"


# ---------- Timetable ----------
@app.route('/timetable')
def timetable():
    if 'username' not in session:
        return redirect(url_for('login'))
    return render_template('timetable.html', username=session['username'])

@app.route('/save-timetable', methods=['POST'])
def save_timetable():
    if 'username' not in session:
        return jsonify({'error': 'Not logged in'}), 401
    data = request.json
    if not data or 'timetable' not in data:
        return jsonify({'error': 'Invalid data'}), 400
    db = load_json('timetable.json', {})
    db[session['username']] = data['timetable']
    save_json('timetable.json', db)
    return jsonify({'message': 'Timetable saved successfully'})

@app.route('/load-timetable')
def load_timetable():
    if 'username' not in session:
        return jsonify({'error': 'Not logged in'}), 401
    db = load_json('timetable.json', {})
    return jsonify(db.get(session['username'], {}))

@app.route('/get-timetable')
def get_timetable():
    return load_timetable()

# ---------- Subject Management ----------
@app.route('/update-subjects', methods=['POST'])
def update_subjects():
    if 'username' not in session:
        return jsonify({'error': 'Not logged in'}), 401
    data = request.json or {}
    subjects = data.get('subjects', [])
    db = load_json('subjects.json', {})
    db[session['username']] = subjects
    save_json('subjects.json', db)
    return jsonify({'message': 'Subjects updated', 'subjects': subjects})

@app.route('/get-subjects')
def get_subjects():
    if 'username' not in session:
        return jsonify([])
    db = load_json('subjects.json', {})
    return jsonify(db.get(session['username'], []))

# ---------- Chatbot ----------
def preprocess(text):
    doc = nlp(text.lower())
    return [token.lemma_ for token in doc if not token.is_stop and not token.is_punct]

INTENTS = {
    "greeting": ["hi", "hello", "hey", "greetings"],
    "goodbye": ["bye", "goodbye", "see you", "farewell"],
    "thanks": ["thanks", "thank you", "appreciate"],
    "schedule": ["schedule", "timetable", "class time", "routine"],
    "reminder": ["reminder", "alert", "note", "remember"],
    "upload": ["upload", "files", "pdf", "documents"],
    "assignment": ["assignment", "homework", "task"],
}

RESPONSES = {
    "greeting": ["Hey! 👋 What can I help you with?", "Hello! How can I assist you today?"],
    "goodbye": ["Goodbye! 👋", "See you soon!"],
    "thanks": ["You're welcome! 🙂", "Happy to help!"],
    "schedule": ["You can view your timetable in the Timetable section.", "Check your schedule page for class timings."],
    "reminder": ["Add reminders from the Reminders page.", "Set alerts to never miss important events."],
    "upload": ["Upload files under the Upload tab.", "You can upload PDFs and documents easily."],
    "assignment": ["Track your assignments in the Assignments section.", "Check your tasks and deadlines here."],
    "default": ["I'm not sure I understand. Could you rephrase?", "Sorry, I didn't get that. Please ask differently."]
}

@app.route('/chatbot', methods=['POST'])
def chatbot():
    user_msg = request.json.get('message', '').strip()
    if not user_msg:
        return jsonify({"response": "Please enter a message."})
    tokens = preprocess(user_msg)
    intent_scores = {}
    for intent, keywords in INTENTS.items():
        intent_tokens = [kw.lower() for kw in keywords]
        score = len(set(tokens) & set(intent_tokens))
        intent_scores[intent] = score
    best_intent = max(intent_scores, key=intent_scores.get)
    if intent_scores[best_intent] == 0:
        best_intent = "default"
    reply = random.choice(RESPONSES.get(best_intent, RESPONSES["default"]))
    return jsonify({"response": reply})

# ---------- Spotify ----------
@app.route('/save-spotify', methods=['POST'])
def save_spotify():
    username = request.form.get('username') or session.get('username')
    if not username:
        return redirect('/')
    raw = request.form.get('spotify_url', '').strip()
    embed = to_spotify_embed(raw) if raw else "https://open.spotify.com/embed/playlist/37i9dQZF1DXcBWIGoYBM5M"
    db = load_json('spotify.json', {})
    db[username] = embed
    save_json('spotify.json', db)
    return redirect(f"/dashboard?user={username}")

# ------------------------
# Notes with JSON persistence
# ------------------------
@app.route('/notes', methods=['GET', 'POST'])
def notes_page():
    user = request.args.get('user')
    if not user:
        return "User not specified", 400
    notes = load_notes(user)
    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'save':
            title = request.form.get('title', '').strip()
            content = request.form.get('content', '').strip()
            date_str = request.form.get('date', '').strip()
            index = request.form.get('index')
            if not title:
                return "Title is required", 400
            if not date_str:
                date_str = datetime.today().strftime('%Y-%m-%d')
            note = {
                'title': title,
                'content': content,
                'date': date_str
            }
            if index and index.isdigit():
                idx = int(index)
                if 0 <= idx < len(notes):
                    notes[idx] = note
                else:
                    return "Invalid note index", 400
            else:
                notes.append(note)
            save_notes(user, notes)
            return redirect(url_for('notes_page', user=user))
        elif action == 'delete':
            index = request.form.get('index')
            if index and index.isdigit():
                idx = int(index)
                if 0 <= idx < len(notes):
                    notes.pop(idx)
                    save_notes(user, notes)
                    return redirect(url_for('notes_page', user=user))
                else:
                    return "Invalid note index", 400
    return render_template('notes.html', user=user, notes=notes)

# ---------- Main ----------
if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0', port=5000)
